import streamlit as st
from docx import Document
from datetime import datetime
from io import BytesIO

# Función para reemplazar los marcadores de posición en el documento
def replace_placeholder(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholder(cell, placeholder, replacement)

# Función para generar el documento de Word
def create_document(data, template_bytes):
    doc = Document(BytesIO(template_bytes))
    for placeholder, replacement in data.items():
        replace_placeholder(doc, placeholder, replacement)
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Título del formulario
st.title("Formulario Único Desconocimiento de Transacciones")

# Inicializar el estado de sesión si no existe
if 'show_text' not in st.session_state:
    st.session_state.show_text = False

# Función para alternar la visibilidad del texto
def toggle_text():
    st.session_state.show_text = not st.session_state.show_text

# Botón para alternar la visibilidad del texto
if st.button("Script"):
    toggle_text()

# Mostrar u ocultar el texto según el estado
if st.session_state.show_text:
    st.write("""
        Don / Sra.…. de conformidad a lo dispuesto en la Ley 20.009, le informo que, 
        para hacer efectiva su reclamación sobre aquellas operaciones respecto de las 
        cuales desconoce haber otorgado su autorización o consentimiento, se solicita 
        completar en su totalidad la Declaración Jurada que le enviaremos al correo 
        electrónico que tiene registrado en nuestros sistemas y enviarla junto a su 
        Cédula de Identidad, en un plazo de 30 días hábiles. 

        En esa misma instancia, y de acuerdo a lo señalado por la Ley, se requiere para 
        la cancelación de los cargos o la restitución de los fondos correspondientes a 
        las operaciones reclamadas, que nos haga llegar también el respaldo de la 
        denuncia por los hechos constitutivos del delito de fraude que usted debe 
        efectuar ante al menos uno de los siguientes organismos: el Ministerio Público, 
        funcionarios de Carabineros de Chile, de la Policía de Investigaciones, o ante 
        cualquier tribunal con competencia criminal. 

        Le sugerimos consultar con sus familiares o adicional si alguno realizó estas 
        transacciones que usted está desconociendo. 

        Finalmente es necesario indicarle, que se realizará una investigación interna de 
        los desconocimientos realizados por los clientes, y en la medida que se cuente 
        con evidencias de mecanismos robustos de seguridad, la empresa podrá hacer uso 
        de las acciones legales que otorga la ley.
    """)


st.header("I. Antecedentes personales")
nombre = st.text_input("Nombre Tarjetahabiente (Cardholder name)")
tc = st.text_input("N° tarjeta (Cardholder number) 4 últimos dígitos")
tt_tarjeta = st.selectbox("Tipo de tarjeta (Card type)", ["Titular", "Adicional", "Ambas"])
direccion = st.text_input("Dirección (Address)")
correo = st.text_input("E-Mail")
telefono = st.text_input("Celular (Cel Phone)")
rut = st.text_input("RUT")

# Fecha de reclamo
fecha_actual = datetime.now().strftime("%d/%m/%Y")

# Detalle de transacciones reclamadas
st.header("III. Detalle Transacciones Reclamadas (Transaction Details)")
moneda = st.selectbox("Moneda para la suma total", ["Pesos", "Dólares"])
monto_total_label = "USD" if moneda == "Dólares" else "$"
    
num_transacciones = st.number_input("Cantidad de transacciones reclamadas", min_value=1, max_value=15, step=1)
transacciones = []
monto_total = 0.0  # Variable para sumar los montos

for i in range(num_transacciones):
    fecha = st.text_input(f"Fecha (dd/mm/aa) - Transacción {i+1}")
    nombre_comercio = st.text_input(f"Nombre del Comercio - Transacción {i+1}")
    monto = st.number_input(f"Monto - Transacción {i+1}", min_value=0.0, format="%.2f")
    transacciones.append((fecha, nombre_comercio, monto))
    monto_total += monto  # Sumar el monto a la variable de suma total

# Observaciones
st.header("IV. Observations (Observaciones)")
observaciones = st.text_area("OBSERVACIONES")

def mostrar_script():
    st.code("""
    # Aquí va tu script que deseas mostrar
    print('Don / Sra.…. de conformidad a lo dispuesto en la Ley 20.009')
    """, language='python')



# Generar el documento automáticamente
if nombre and tc and direccion and correo and telefono and rut and num_transacciones > 0:
    # Ruta al archivo de plantilla en el repositorio
    template_path = "static/datasets/template.docx"

    try:
        with open(template_path, "rb") as template_file:
            template_bytes = template_file.read()

        data = {
            "{{Nombre}}": nombre,
            "{{Tc}}": tc,
            "Titular": tt_tarjeta,
            "Direccion": direccion,
            "{{Correo}}": correo,
            "{{Teléfono}}": telefono,
            "{{Rut}}": rut,
            "Numero TRX": str(num_transacciones),
            "{{Run}}": f"{monto_total_label} {monto_total:.2f}",  # Concatenar la moneda y el monto total
            "{{Observación}}": observaciones,
            "Input_Observaciones": observaciones,
            "Fecha actual": fecha_actual,
            "Fecha-5": "",
            "Fecha-4": "",
            "Fecha-3": "",
            "Fecha-2": "",
            "Fecha-1": "",
            "Fecha0": "",
            "Fecha1": "",
            "Fecha2": "",
            "Fecha3": "",
            "Fecha4": "",
            "Fecha5": "",
            "Fecha6": "",
            "Fecha7": "",
            "Fecha8": "",
            "Fecha9": "",
            "NombreComercio-5": "",
            "NombreComercio-4": "",
            "NombreComercio-3": "",
            "NombreComercio-2": "",
            "NombreComercio-1": "",
            "NombreComercio0": "",
            "NombreComercio1": "",
            "NombreComercio2": "",
            "NombreComercio3": "",
            "NombreComercio4": "",
            "NombreComercio5": "",
            "NombreComercio6": "",
            "NombreComercio7": "",
            "NombreComercio8": "",
            "NombreComercio9": "",
            "Monto-5": "",
            "Monto-4": "",
            "Monto-3": "",
            "Monto-2": "",
            "Monto-1": "",
            "Monto0": "",
            "Monto1": "",
            "Monto2": "",
            "Monto3": "",
            "Monto4": "",
            "Monto5": "",
            "Monto6": "",
            "Monto7": "",
            "Monto8": "",
            "Monto9": ""
            
        }

        for a, (fecha, nombre_comercio, monto) in enumerate(transacciones, start=-5):
            data[f"Fecha{a}"] = fecha
            data[f"NombreComercio{a}"] = nombre_comercio
            data[f"Monto{a}"] = f"{monto_total_label} {monto:.2f}"
        doc_file = create_document(data, template_bytes)
        
        st.success("Documento actualizado y listo para descargar.")
        st.download_button(label="Descargar Documento", data=doc_file, file_name="Formulario_unico_desconocimiento.docx")
        st.write(f"Cliente desconoce {num_transacciones} compras según detalle de archivo. Correo: {correo} / Teléfono: {telefono}")

    except FileNotFoundError:
        st.error("No se encontró la plantilla de Word en el repositorio.")
    except Exception as e:
        st.error(f"Error al generar el documento: {e}")
else:
    st.info("Por favor, complete todos los campos requeridos para generar el documento.")
